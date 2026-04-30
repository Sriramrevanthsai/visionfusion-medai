from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "VisionFusion_MedAI_Project_Report.docx"


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.add_run(text)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        run = hdr[i].paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        hdr[i]._tc.get_or_add_tcPr().append(parse_shading("0B7A75"))
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()
    return table


def parse_shading(color: str):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    return parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
for name, size, color in [("Heading 1", 18, "075B58"), ("Heading 2", 14, "0B7A75"), ("Heading 3", 12, "15211F")]:
    style = styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("VisionFusion MedAI\n")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(7, 91, 88)
sub = title.add_run("Voice, Vision, and Text Intelligence for Healthcare")
sub.font.size = Pt(15)
sub.font.color.rgb = RGBColor(96, 112, 108)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("\nProject Report\n").bold = True
meta.add_run("Technology Stack: FastAPI, MedGemma, MedASR, MCP, Python, HTML/CSS/JavaScript\n")
meta.add_run("Sector: Healthcare\n")
meta.add_run("Prepared for academic project submission")

doc.add_page_break()

add_heading(doc, "Project Description")
doc.add_paragraph(
    "VisionFusion MedAI is a multimodal healthcare insight assistant designed to support patients, doctors, clinics, and hospital teams through AI-powered interpretation of medical images, voice notes, and text-based symptoms. The system combines medical image comprehension, medical speech transcription, structured clinical reasoning, risk triage, and automated report generation within a polished web interface."
)
doc.add_paragraph(
    "The platform uses MedGemma for medical image and text comprehension, MedASR for medical dictation and physician-patient conversation transcription, and an MCP-style tool layer to connect specialized tools with LLM reasoning."
)

add_heading(doc, "Architecture Overview")
for item in [
    "User logs in or registers as patient, doctor, or admin.",
    "User uploads a medical image, optional voice note, and typed symptoms.",
    "FastAPI stores files under a unique case ID.",
    "MCP tools run image feature extraction, audio transcription, risk scoring, and report generation.",
    "MedGemma-style reasoning combines image, transcript, and symptom context.",
    "The user receives doctor-mode or patient-mode output and downloads a PDF report.",
]:
    doc.add_paragraph(item, style="List Bullet")

add_heading(doc, "Core Technologies")
add_table(
    doc,
    ["Technology", "Purpose"],
    [
        ["FastAPI", "Backend APIs, uploads, authentication, analysis routes"],
        ["MedGemma", "Medical image and text reasoning"],
        ["MedASR", "Medical speech-to-text for dictation and conversations"],
        ["MCP Tool Layer", "Coordinates image, audio, risk, retrieval, and report tools"],
        ["HTML/CSS/JavaScript", "High-standard responsive user interface"],
        ["PDF Generator", "Downloadable healthcare case reports"],
        ["SQLite/PostgreSQL", "Production persistence for users, cases, and audit logs"],
    ],
)

doc.add_page_break()
add_heading(doc, "Key Functionalities")
add_table(
    doc,
    ["No.", "Functionality", "Description"],
    [
        ["1", "Login and Registration", "Role-aware access for doctors, patients, and admins."],
        ["2", "Medical Image Upload", "Users upload health-related images for AI-assisted interpretation."],
        ["3", "Voice Note Upload", "Doctors can upload dictation or patient conversation audio."],
        ["4", "MedASR Transcription", "Medical speech is converted into text for downstream reasoning."],
        ["5", "Symptom + Image Reasoning", "Typed symptoms are combined with image and voice context."],
        ["6", "Doctor Mode", "Structured clinical support summary and follow-up questions."],
        ["7", "Patient Mode", "Simple explanation and safe next steps."],
        ["8", "Risk Triage Engine", "Low, moderate, or urgent risk classification."],
        ["9", "Case History", "Recent case insights are displayed in the dashboard."],
        ["10", "PDF Report Download", "Users can download a structured case report."],
    ],
)

add_heading(doc, "Important Use Cases")
add_table(
    doc,
    ["Use Case", "Description", "Primary User"],
    [
        ["Remote Patient Pre-Screening", "Patient uploads image and symptoms for safe next-step guidance.", "Patient"],
        ["Doctor Dictation Support", "MedASR transcribes doctor notes and prepares structured summaries.", "Doctor"],
        ["Wound or Skin Follow-Up", "Repeat images help summarize visible improvement or worsening.", "Doctor/Patient"],
        ["Report Generation", "System creates PDF reports for consultation or records.", "Doctor/Patient"],
        ["Clinic Workflow Assistant", "Clinics organize image, voice, and text before doctor review.", "Clinic"],
        ["Research Evaluation", "Compare text-only, image+text, and image+voice+text reasoning.", "Researcher"],
    ],
)

doc.add_page_break()
add_heading(doc, "API Endpoints")
add_table(
    doc,
    ["Endpoint", "Method", "Purpose"],
    [
        ["/api/auth/register", "POST", "Create new user"],
        ["/api/auth/login", "POST", "Authenticate existing user"],
        ["/api/demo-user", "GET", "Open demo doctor workspace"],
        ["/api/cases/analyze", "POST", "Generate multimodal insight"],
        ["/api/cases", "GET", "Retrieve case history"],
        ["/api/cases/{case_id}/report", "GET", "Download PDF report"],
    ],
)

add_heading(doc, "Research-Level Enhancements")
for item in [
    "Compare MedGemma 4B and 27B multimodal outputs.",
    "Compare MedASR with general ASR for medical terminology transcription.",
    "Evaluate text-only, image+text, and image+voice+text reasoning quality.",
    "Add doctor feedback scoring and report correction tracking.",
    "Add retrieval-augmented generation from verified medical references.",
    "Track latency, risk classification quality, and user satisfaction.",
]:
    doc.add_paragraph(item, style="List Bullet")

add_heading(doc, "Market-Ready Features")
add_table(
    doc,
    ["Feature", "Market Value"],
    [
        ["Role-Based Access", "Supports doctors, patients, clinic staff, and admins"],
        ["Human Review Workflow", "Keeps clinicians in control of final decisions"],
        ["PDF Export", "Easy sharing and archiving"],
        ["Secure Storage Plan", "Required for healthcare deployment"],
        ["FHIR/HL7 Readiness", "Future hospital integration"],
        ["Audit Logs", "Compliance and accountability"],
        ["Cloud/Local Model Options", "Supports privacy-first and scalable deployments"],
    ],
)

add_heading(doc, "Safety and Ethical Considerations")
doc.add_paragraph(
    "VisionFusion MedAI is not a replacement for doctors and must not be presented as an automatic diagnosis system. It is a clinical decision-support and patient education assistant. Production deployment requires encryption, consent, access control, audit logs, bias evaluation, human review, and urgent-case escalation."
)

doc.add_page_break()
add_heading(doc, "Conclusion")
doc.add_paragraph(
    "VisionFusion MedAI demonstrates a research-level and market-ready direction for healthcare AI by combining image understanding, medical speech transcription, text reasoning, MCP tool orchestration, and professional report generation. The project goes beyond a basic multimodal demo by including authentication, role-aware workflows, risk triage, case history, and downloadable reports."
)

doc.save(OUT)
print(OUT)
